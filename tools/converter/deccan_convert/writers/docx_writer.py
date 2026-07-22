"""DOCX writer: DocumentIR -> .docx built on the Deccan Word template.

The bundled deccan-document-base.docx is templates/word/deccan-document.dotx
re-zipped as a document (see sync_assets.py), so its styles, theme, page
setup, three-section structure (cover / body / end page), and the body
footer ("Deccan Fine Chemicals · Confidential" + bare page number, Cascadia
Mono 8.5pt) are the .dotx's verbatim. This writer fills the cover
placeholders, replaces the sample body with the IR content, and stamps the
classification — fonts are referenced by name only, never embedded.
"""

from __future__ import annotations

import base64
import io
import re
from pathlib import Path

import docx
from bs4 import BeautifulSoup, NavigableString, Tag
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from deccan_convert.assets import asset_path
from deccan_convert.ir import DocumentIR

# Design tokens (skill/references/tokens.md) needed for direct formatting
# where the template styles don't reach (table cells, hyperlinks).
DECCAN_BLUE = RGBColor(0x16, 0x49, 0x99)
STONE_100 = "F5F5F4"
STONE_200 = "E7E5E4"
STONE_50 = "FAFAF9"
STONE_700 = RGBColor(0x57, 0x53, 0x4E)
STONE_900 = RGBColor(0x1C, 0x19, 0x17)
MONO_FONT = "Cascadia Mono"

_COVER_PLACEHOLDERS = {
    "DOCUMENT": lambda m: m.document_type.upper(),
    "Document title": lambda m: m.title,
    "One-sentence subtitle": lambda m: m.subtitle,
}

_META_TABLE_ORDER = ("document_type", "prepared_by", "date", "version", "classification")


def write_docx(ir: DocumentIR, path: Path) -> Path:
    meta = ir.metadata.with_defaults()
    missing = meta.missing_required()
    if missing:
        raise ValueError(
            "Missing required document details: " + ", ".join(missing) + ". "
            "These are never invented — provide them explicitly."
        )

    doc = docx.Document(str(asset_path("deccan-document-base.docx")))

    _fill_cover(doc, meta)
    _fill_end_page(doc, meta)
    anchor = _clear_sample_body(doc)
    _build_body(doc, anchor, ir)
    _set_core_properties(doc, meta)

    doc.save(str(path))
    return path


# --- template plumbing -------------------------------------------------------


def _section_break_paragraphs(doc) -> list:
    """Paragraphs that carry a w:sectPr (they terminate sections 0 and 1)."""
    found = []
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            found.append(p)
    return found


def _replace_paragraph_text(paragraph, text: str) -> None:
    """Set text while keeping the first run's formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _index_of(paragraphs, target) -> int:
    for i, p in enumerate(paragraphs):
        if p._p is target._p:
            return i
    raise ValueError("paragraph not found in document body")


def _fill_cover(doc, meta) -> None:
    breaks = _section_break_paragraphs(doc)
    paragraphs = doc.paragraphs
    cover_end = _index_of(paragraphs, breaks[0]) if breaks else len(paragraphs)
    for p in paragraphs[:cover_end]:
        replacer = _COVER_PLACEHOLDERS.get(p.text.strip())
        if replacer is not None:
            _replace_paragraph_text(p, replacer(meta))

    # Cover metadata strip: a 2-row table, labels then values.
    if doc.tables:
        table = doc.tables[0]
        if len(table.rows) >= 2:
            values = [getattr(meta, field) for field in _META_TABLE_ORDER]
            for cell, value in zip(table.rows[1].cells, values):
                _replace_paragraph_text(cell.paragraphs[0], value)


def _fill_end_page(doc, meta) -> None:
    for p in doc.paragraphs:
        if "INTERNAL USE" in p.text.upper() and "·" in p.text:
            _replace_paragraph_text(p, f"{meta.classification.upper()} · INTERNAL USE")


def _clear_sample_body(doc):
    """Remove the template's sample body paragraphs; return the section-break
    paragraph new content must be inserted before."""
    breaks = _section_break_paragraphs(doc)
    if len(breaks) < 2:
        raise RuntimeError("Base docx does not have the expected 3-section layout")
    paragraphs = doc.paragraphs
    start = _index_of(paragraphs, breaks[0]) + 1
    end = _index_of(paragraphs, breaks[1])
    for p in paragraphs[start:end]:
        p._p.getparent().remove(p._p)
    return breaks[1]


# --- body construction -------------------------------------------------------


class _BodyBuilder:
    """Builds docx content from the IR body HTML, inserted before `anchor`."""

    def __init__(self, doc, anchor, ir: DocumentIR):
        self.doc = doc
        self.anchor = anchor._p
        self.ir = ir
        self.first_h1_seen = False
        self.style_names = {s.name for s in doc.styles}

    # Every python-docx add_* call appends at the document end (inside the
    # end-page section); _move relocates the element before the anchor.
    def _move(self, element) -> None:
        self.anchor.addprevious(element)

    def _style(self, wanted: str, fallback: str = "Normal") -> str:
        return wanted if wanted in self.style_names else fallback

    def add(self, node: Tag) -> None:
        name = node.name
        if name == "h1":
            self._heading(node, 1)
        elif name in ("h2", "h3", "h4"):
            self._heading(node, int(name[1]))
        elif name == "p":
            classes = node.get("class", ())
            style = "Lead" if "lead" in classes else "Normal"
            self._paragraph(node, style=self._style(style))
        elif name in ("ul", "ol"):
            self._list(node, ordered=(name == "ol"), level=0)
        elif name == "table":
            self._table(node)
        elif name == "pre":
            self._code_block(node)
        elif name == "div":
            classes = node.get("class", ())
            if "callout" in classes:
                self._callout(node, muted="muted" in classes)
            elif "pullquote" in classes:
                self._pullquote(node)
            else:
                for child in node.find_all(recursive=False):
                    self.add(child)
        elif name in ("dl",):
            self._definition_list(node)
        elif name == "hr":
            p = self.doc.add_paragraph()
            self._move(p._p)
        elif name in ("img",):
            self._image(node)
        else:
            self._paragraph(node)

    def _heading(self, node: Tag, level: int) -> None:
        num_span = node.find("span", class_="num")
        num_text = ""
        if num_span is not None:
            num_text = num_span.get_text(strip=True)
            num_span.decompose()
        text = node.get_text(strip=True)
        if level == 1 and num_text:
            text = f"{num_text}  {text}"
        p = self.doc.add_paragraph(style=self._style(f"Heading {level}"))
        p.add_run(text)
        if level == 1:
            # Every H1 on a new page except the first body section.
            p.paragraph_format.page_break_before = self.first_h1_seen
            self.first_h1_seen = True
        self._move(p._p)

    def _paragraph(self, node: Tag, style: str = "Normal") -> None:
        if not node.get_text(strip=True) and not node.find("img"):
            return
        p = self.doc.add_paragraph(style=style)
        self._inline(p, node)
        self._move(p._p)

    def _list(self, node: Tag, ordered: bool, level: int) -> None:
        base = "List Number" if ordered else "List Bullet"
        suffix = "" if level == 0 else f" {min(level + 1, 3)}"
        style = self._style(base + suffix, fallback=self._style(base))
        for li in node.find_all("li", recursive=False):
            nested = li.find_all(["ul", "ol"], recursive=False)
            for sub in nested:
                sub.extract()
            p = self.doc.add_paragraph(style=style)
            self._inline(p, li)
            self._move(p._p)
            for sub in nested:
                self._list(sub, ordered=(sub.name == "ol"), level=level + 1)

    def _code_block(self, node: Tag) -> None:
        p = self.doc.add_paragraph(style=self._style("Code Block"))
        run = p.add_run(node.get_text().rstrip("\n"))
        if "Code Block" not in self.style_names:
            run.font.name = MONO_FONT
            run.font.size = Pt(9)
            _shade_paragraph(p, STONE_100)
        self._move(p._p)

    def _callout(self, node: Tag, muted: bool) -> None:
        style = "Callout Muted" if muted else "Callout Default"
        label = node.find(class_="label")
        if label is not None:
            label_p = self.doc.add_paragraph(style=self._style(style))
            run = label_p.add_run(label.get_text(strip=True).upper())
            run.font.name = MONO_FONT
            run.font.size = Pt(8.5)
            run.font.bold = True
            self._move(label_p._p)
            label.decompose()
        paragraphs = node.find_all("p", recursive=False) or [node]
        for para in paragraphs:
            p = self.doc.add_paragraph(style=self._style(style))
            self._inline(p, para)
            self._move(p._p)

    def _pullquote(self, node: Tag) -> None:
        p = self.doc.add_paragraph(style=self._style("Intense Quote", "Quote"))
        self._inline(p, node)
        self._move(p._p)

    def _definition_list(self, node: Tag) -> None:
        for child in node.find_all(["dt", "dd"], recursive=False):
            if child.name == "dt":
                p = self.doc.add_paragraph()
                run = p.add_run(child.get_text(strip=True))
                run.font.bold = True
            else:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                self._inline(p, child)
            self._move(p._p)

    def _table(self, node: Tag) -> None:
        rows = node.find_all("tr")
        if not rows:
            return
        n_cols = max(len(tr.find_all(["th", "td"])) for tr in rows)
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        table.style = self.doc.styles["Table Grid"] if "Table Grid" in self.style_names else None
        header_rows = {id(tr) for tr in rows if tr.find_parent("thead") is not None or tr.find("th")}
        for r, tr in enumerate(rows):
            cells = tr.find_all(["th", "td"])
            is_header = id(tr) in header_rows
            for c in range(n_cols):
                cell = table.rows[r].cells[c]
                text = cells[c].get_text(" ", strip=True) if c < len(cells) else ""
                para = cell.paragraphs[0]
                run = para.add_run(text.upper() if is_header else text)
                if is_header:
                    # thead contract: mono caps, stone-700 on stone-100.
                    run.font.name = MONO_FONT
                    run.font.size = Pt(8.5)
                    run.font.bold = True
                    run.font.color.rgb = STONE_700
                    _shade_cell(cell, STONE_100)
                else:
                    run.font.size = Pt(9.5)
                    if r % 2 == 0:
                        _shade_cell(cell, STONE_50)
        self._move(table._tbl)
        spacer = self.doc.add_paragraph()
        self._move(spacer._p)

    def _image(self, node: Tag) -> None:
        src = node.get("src", "")
        if src.startswith("data:image/"):
            try:
                payload = src.split(",", 1)[1]
                stream = io.BytesIO(base64.b64decode(payload))
                p = self.doc.add_paragraph()
                run = p.add_run()
                run.add_picture(stream, width=Inches(5.3))
                self._move(p._p)
                return
            except Exception:
                pass
        self.ir.warnings.append(
            f"docx: image '{src[:60]}' could not be embedded and was skipped."
        )

    # --- inline content ---

    def _inline(self, paragraph, node: Tag) -> None:
        for child in node.children:
            self._inline_node(paragraph, child, bold=False, italic=False, code=False)

    def _inline_node(self, paragraph, child, bold: bool, italic: bool, code: bool) -> None:
        if isinstance(child, NavigableString):
            text = re.sub(r"\s+", " ", str(child))
            if not text:
                return
            run = paragraph.add_run(text)
            run.font.bold = bold or None
            run.font.italic = italic or None
            if code:
                _apply_inline_code(run, self.style_names, paragraph)
            return
        if not isinstance(child, Tag):
            return
        if child.name == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
            return
        if child.name == "img":
            self._inline_image(paragraph, child)
            return
        if child.name == "a":
            self._hyperlink(paragraph, child)
            return
        bold = bold or child.name == "strong"
        italic = italic or child.name == "em"
        code = code or child.name == "code"
        for grandchild in child.children:
            self._inline_node(paragraph, grandchild, bold, italic, code)

    def _inline_image(self, paragraph, node: Tag) -> None:
        src = node.get("src", "")
        if src.startswith("data:image/"):
            try:
                payload = src.split(",", 1)[1]
                stream = io.BytesIO(base64.b64decode(payload))
                paragraph.add_run().add_picture(stream, width=Inches(5.3))
                return
            except Exception:
                pass
        self.ir.warnings.append(
            f"docx: image '{src[:60]}' could not be embedded and was skipped."
        )

    def _hyperlink(self, paragraph, node: Tag) -> None:
        url = node.get("href", "")
        text = node.get_text(" ", strip=True) or url
        if not url:
            run = paragraph.add_run(text)
            return
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "164999")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(color)
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)


def _apply_inline_code(run, style_names, paragraph) -> None:
    if "Code Inline" in style_names:
        run.style = "Code Inline"
    else:
        run.font.name = MONO_FONT
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), STONE_100)
        rPr.append(shd)


def _shade_cell(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _shade_paragraph(paragraph, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def _build_body(doc, anchor, ir: DocumentIR) -> None:
    soup = BeautifulSoup(ir.body_html, "html.parser")
    builder = _BodyBuilder(doc, anchor, ir)
    sections = soup.find_all("section", class_="section") or [soup]
    for section in sections:
        for child in section.find_all(recursive=False):
            builder.add(child)


def _set_core_properties(doc, meta) -> None:
    props = doc.core_properties
    props.title = meta.title
    props.subject = meta.subtitle
    props.author = meta.prepared_by
    props.category = meta.document_type
    props.comments = f"deccan-design v2.0 · {meta.classification}"
